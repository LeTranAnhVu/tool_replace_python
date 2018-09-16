from docx import Document
from error_handler import ErrorHandler

DEFAULT_PATH = r'.\sample\ex1.docx'


def open_docx(path):
    return Document(path)





def replace_docx(path, old_str, new_str):
    # open docx
    doc = open_docx(path)
    # check none
    if not isinstance(old_str, str) or not isinstance(new_str, str):
        e = "the parameter in replace_docx must not be None"
        return ErrorHandler(path, e)

    # replace the string
    try:
        for i, para in enumerate(doc.paragraphs):
            # para.text = para.text.replace(old_str,new_str)

            # print(para.text)
            # para.text = para.text.replace(old_str, new_str)
            for j, run in enumerate(para.runs):
                print(">>:", j)
                print(run.text)
                print("<<")
            # print(">>:", i)
            # print( para.text)
            # print(para.style.quick_style)
            # print("<<")
            # print("write Ok")
        # doc.save(DEFAULT_PATH)
    except AttributeError as e:
        errorObj = ErrorHandler(path, e)
        errorObj.showError()
        return errorObj

    # save the file


# replace_docx(DEFAULT_PATH, 'ge', newStr)
# word = 'geeks for geeks sjdhfshgedfhsgefsdj;fjasjdfj;age hadsjlfjhaldhflage ;jsdkjjhalkjdhgege'
# print(word.find('ge'))


def matchIndexs(str, offset, pattern):
    itemList = []
    i = str.find(pattern, offset)
    length = len(str)
    patternLen = len(pattern)
    while 0 <= i <= length:
        itemList.append(i)
        i = str.find(pattern, i + patternLen)
    return itemList


oldStr = "trang"
newStr = r"Tran Thi Ut"
doct = Document(DEFAULT_PATH)

lastGapsOfWholeDoc = []
firstGapsOfWholeDoc = []


def getRawRun(doc, oldStr):
    listOfNeedRunIndicesInAllDoc = []
    for ipara, p in enumerate(doc.paragraphs):
        if oldStr in p.text:
            indexs = matchIndexs(p.text, 0, oldStr)
            print('indexs: ', indexs)
            firstGaps = []
            lastGaps = []
            for index in indexs:
                indexOfWholeStartRun = -1
                indexOfStartRun = 0
                if indexOfWholeStartRun < index:
                    for irun, run in enumerate(p.runs):
                        indexOfWholeStartRun += len(run.text)
                        if indexOfWholeStartRun >= index:
                            indexOfStartRun = irun
                            break
                print('indexOfWholeStartRun: ', indexOfWholeStartRun)
                firstGap = indexOfWholeStartRun - index
                firstGaps.append(firstGap)

                print('firstGap:', firstGap)
                # get the start run
                startRun = p.runs[indexOfStartRun]
                indexOfFirstCharInStartRun = len(startRun.text) - firstGap - 1
                print('indexOfFirstCharInStartRun: ', indexOfFirstCharInStartRun)
                print("result: ", startRun.text[indexOfFirstCharInStartRun:])
                # get the stop run
                lenOfOldStr = len(oldStr)
                templen = len(startRun.text[indexOfFirstCharInStartRun:])
                nextRun = indexOfStartRun + 1
                indexOfStopRun = indexOfStartRun
                listOfNeedRunIndices = []
                listOfNeedRunIndices.append(indexOfStartRun)
                lastGap = 0
                if templen < lenOfOldStr:
                    for irun, run in enumerate(p.runs[nextRun:]):
                        templen += len(run.text)
                        if templen >= lenOfOldStr:
                            indexOfStopRun = irun + nextRun
                            lastGap = templen - lenOfOldStr
                            listOfNeedRunIndices.append(indexOfStopRun)
                            break
                        else:
                            listOfNeedRunIndices.append(irun + nextRun)
                stopRun = p.runs[indexOfStopRun]
                print('stopRun.text: |' + stopRun.text + "|")
                print('lastGap: ', lastGap)
                print('listOfNeedRunIndices:', listOfNeedRunIndices)
                lastGaps.append(lastGap)
                # tempDict = {}
                # tempDict[ipara]= listOfNeedRunIndices
                listOfNeedRunIndicesInAllDoc.append(listOfNeedRunIndices)
            firstGapsOfWholeDoc.append(firstGaps)
            lastGapsOfWholeDoc.append(lastGaps)
    return listOfNeedRunIndicesInAllDoc


listOfNeedRunIndicesInAllDoc = getRawRun(doct, oldStr)
print("listOfNeedRunIndicesInAllDoc: ", listOfNeedRunIndicesInAllDoc)
print("firstGapsOfWholeDoc: ", firstGapsOfWholeDoc)
print("lastGapsOfWholeDoc: ", lastGapsOfWholeDoc)

lastGap = lastGapsOfWholeDoc[0][0]
firstGap = firstGapsOfWholeDoc[0][0]
print("<<<<lastGap:", lastGap)
print("<<<<firstGap:", firstGap)


def replaceString(doc, old_str, new_str, listOfNeedRunIndices, firstGap, lastGap):
    # delete all the middle runs
    runs = doc.paragraphs[-1].runs
    if len(listOfNeedRunIndices) > 1 :
        for irun, run in enumerate(runs):
            print(">>>>",irun)
            print(run.text)
            print("<<<<")
        for i in listOfNeedRunIndices[1:-1]:
            print(">>>>>i:", i)
            runs[i].text = ""
        print('hey:','|'+runs[1].text+'|')
        # filter the last run
        lenOfStringInLastRun = len(runs[listOfNeedRunIndices[-1]].text)
        print("lenOfStringInLastRun: ", lenOfStringInLastRun)
        indexOfCharInLastRun = lenOfStringInLastRun - lastGap - 1

        if (indexOfCharInLastRun < 0):
            print("there is some problem in replaceString function, indexOfCharInLastRun < 0 ")
            return None
        print('indexOfCharInLastRun:', indexOfCharInLastRun)
        if lenOfStringInLastRun == 1:
            # special case
            runs[listOfNeedRunIndices[-1]].text = ""
        else:
            runs[listOfNeedRunIndices[-1]].text = runs[listOfNeedRunIndices[-1]].text[indexOfCharInLastRun + 1:]

        # replace the first run

        lenOfStringInFirstRun = len(runs[listOfNeedRunIndices[0]].text)
        print("text: ", runs[1].text)
        indexOfCharInFirstRun = lenOfStringInFirstRun - firstGap - 1
        print('indexOfCharInFirstRun:', indexOfCharInFirstRun)
        if (indexOfCharInFirstRun < 0):
            print("there is some problem in replaceString function : indexOfCharInFirstRun < 0 ")
            return None
        if indexOfCharInFirstRun == 0:
            runs[listOfNeedRunIndices[0]].text = new_str
        else:
            temp = runs[listOfNeedRunIndices[0]].text[:indexOfCharInFirstRun]
            temp += new_str
            runs[listOfNeedRunIndices[0]].text = temp
    else:
        # length of listOfNeedRunIndices is 1, mean that
        lenOfStringInFirstRun = len(runs[listOfNeedRunIndices[0]].text)
        print("text: ", runs[1].text)
        indexOfCharInFirstRun = lenOfStringInFirstRun - firstGap - 1
        print('indexOfCharInFirstRun:', indexOfCharInFirstRun)
        if (indexOfCharInFirstRun < 0):
            print("there is some problem in replaceString function : indexOfCharInFirstRun < 0 ")
            return None
        if indexOfCharInFirstRun == 0:
            runs[listOfNeedRunIndices[0]].text = new_str
        else:
            temp = runs[listOfNeedRunIndices[0]].text[:indexOfCharInFirstRun]
            temp += new_str
            runs[listOfNeedRunIndices[0]].text = temp

    doc.save('.\sample\ex2.docx')




replaceString(doct, oldStr, newStr, listOfNeedRunIndicesInAllDoc[-1], firstGap, lastGap)
