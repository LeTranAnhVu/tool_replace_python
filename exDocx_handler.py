from docx import Document
from error_handler import ErrorHandler

DEFAULT_PATH = r'.\sample\ex1.docx'


class MatchedRunObj:
    def __init__(self, paragraph, matchedIndex, oldStr, newStr):
        self.p = paragraph
        self.matchedIndex = matchedIndex
        self.oldStr = oldStr
        self.newStr = newStr
        self.listOfIndexOfRun = []
        self.indexOfFirstRun = None
        self.indexOfLastRun = None
        # indexOfFirstMatchedChar is the position of the first character-in oldStr in the first Run obj in Run Objlist
        self.indexOfFirstMatchedChar = None
        self.firstOffset = None

        # indexOfLastMatchedChar is the position of the last character-in oldStr in the last Run obj in Run Objlist
        self.indexOfLastMatchedChar = None
        self.lastOffset = None

    def findRelatedRunIndex(self):
        # do something to ...
        lenOldStr = len(self.oldStr)
        matchedIndex = self.matchedIndex
        # print("#####matchedIndex:",matchedIndex )
        # find the first run that contain the matched oldstring
        indexOfCharInWholeRun = -1
        indexOfFirstRun = 0
        if matchedIndex == 0:
            indexOfFirstRun = 0
            indexOfCharInWholeRun += len(self.p.runs[indexOfFirstRun].text)
        elif matchedIndex > 0:
            for i_run, run in enumerate(self.p.runs):
                indexOfCharInWholeRun += len(run.text)
                if indexOfCharInWholeRun >= matchedIndex:
                    # find out the index first run
                    indexOfFirstRun = i_run
                    break
        else:
            print("error, matched index < 0")

        # find the first character position in the first run obj
        print("#####indexOfFirstRun:", indexOfFirstRun)
        indexOfFirstMatchedChar = self.findFirstCharMatchedIndex(indexOfFirstRun, indexOfCharInWholeRun, matchedIndex)[
            0]
        firstRun = self.p.runs[indexOfFirstRun]
        traveledCharLen = len(firstRun.text[indexOfFirstMatchedChar:])
        # find the last run that contain the matched oldstring
        iNextRun = indexOfFirstRun + 1
        indexOfLastRun = indexOfFirstRun
        if traveledCharLen < lenOldStr:
            for i_run, run in enumerate(self.p.runs[iNextRun:]):
                # search from the first matched run
                traveledCharLen += len(run.text)
                if traveledCharLen >= lenOldStr:
                    indexOfLastRun = iNextRun + i_run
                    break
        else:
            pass  # do nothing
        self.findLastCharMatchedIndex(indexOfLastRun, traveledCharLen, lenOldStr)
        self.indexOfFirstRun = indexOfFirstRun
        self.indexOfLastRun = indexOfLastRun
        return self.listOfIndexOfRun

    def findFirstCharMatchedIndex(self, iFirstRun, indexOfCharInWholeRun, matchedIndex):
        # do something to ...
        lenOfFirstRun = len(self.p.runs[iFirstRun].text)
        offset = indexOfCharInWholeRun - matchedIndex
        self.indexOfFirstMatchedChar = lenOfFirstRun - offset - 1
        self.firstOffset = offset
        return [self.indexOfFirstMatchedChar, self.firstOffset]

    def findLastCharMatchedIndex(self, iLastRun, traveledCharLen, lenOldStr):
        # do something to ...
        lenOfLastRun = len(self.p.runs[iLastRun].text)
        offset = traveledCharLen - lenOldStr
        self.indexOfLastMatchedChar = lenOfLastRun - offset - 1
        self.lastOffset = offset
        return [self.indexOfLastMatchedChar, self.lastOffset]

    def replaceString(self):
        # do something to ...
        iFirstRun = self.indexOfFirstRun
        iLastRun = self.indexOfLastRun
        numbOfRun = iLastRun - iFirstRun + 1
        print(">>>>")
        print("iFirstRun: ", iFirstRun)
        print("iLastRun: ", iLastRun)
        print("numbOfRun: ", numbOfRun)
        print("<<<<")
        if numbOfRun == 1:
            # replace to newstring
            self.p.runs[iFirstRun].text = self.p.runs[iFirstRun].text.replace(self.oldStr, self.newStr)
        elif numbOfRun == 2:
            # replace to newstring
            pieceOfStr = self.p.runs[iFirstRun].text[self.indexOfFirstMatchedChar:]
            self.p.runs[iFirstRun].text = self.p.runs[iFirstRun].text.replace(pieceOfStr, self.newStr)

            # remove the rest of oldstring in the 2nd run
            lenOfLastRun = len(self.p.runs[iLastRun].text)
            if (lenOfLastRun - 1) > self.indexOfLastMatchedChar:
                self.p.runs[iLastRun].text = self.p.runs[iLastRun].text[
                                             (self.indexOfLastMatchedChar + 1):]
            elif (lenOfLastRun - 1) == self.indexOfLastMatchedChar:
                self.p.runs[iLastRun].text = ""
            else:
                print("error:(lenOfLastRun - 1) < self.indexOfLastMatchedChar ")

        elif numbOfRun > 2:
            # remove the text of all between runs
            for run in self.p.runs[iFirstRun + 1: iLastRun]:
                run.text = ""
            # replace to newstring
            pieceOfStr = self.p.runs[iFirstRun].text[self.indexOfFirstMatchedChar:]
            self.p.runs[iFirstRun].text = self.p.runs[iFirstRun].text.replace(pieceOfStr, self.newStr)

            # remove the rest of oldstring in the 2nd run
            lenOfLastRun = len(self.p.runs[iLastRun].text)
            if (lenOfLastRun - 1) > self.indexOfLastMatchedChar:
                self.p.runs[iLastRun].text = self.p.runs[iLastRun].text[
                                             (self.indexOfLastMatchedChar + 1):]
            elif (lenOfLastRun - 1) == self.indexOfLastMatchedChar:
                self.p.runs[iLastRun].text = ""
            else:
                print("error:(lenOfLastRun - 1) < self.indexOfLastMatchedChar ")

        else:
            print("error, numbOfRun < 0")
        return True


def findMatchedIndices(str, offset, pattern):
    itemList = []
    i = str.find(pattern, offset)
    length = len(str)
    patternLen = len(pattern)
    while 0 <= i <= length:
        itemList.append(i)
        i = str.find(pattern, i + patternLen)
    return itemList


def open_docx(path):
    return Document(path)


def replace_docx(path, old_str, new_str):
    doc = open_docx(path)
    # check none
    if not isinstance(old_str, str) or not isinstance(new_str, str):
        e = "the parameter in replace_docx must not be None"
        return ErrorHandler(path, e)
    # replace the string
    try:
        # paragraph processing
        for para in doc.paragraphs:
            # each para call  matchIndexs
            if old_str in para.text:
                matchedIndices = findMatchedIndices(para.text, 0, old_str)
                for i_matchIndex, matchIndex in enumerate(matchedIndices):
                    # make the instance of matchedRunObj
                    diff = len(new_str) - len(old_str)
                    updatedMatchedIndex = i_matchIndex * diff + matchIndex
                    matchedObj = MatchedRunObj(para, updatedMatchedIndex, old_str, new_str)
                    matchedObj.findRelatedRunIndex()
                    print("indexOfFirstMatchedChar: ", matchedObj.indexOfFirstMatchedChar)
                    print("indexOfLastMatchedChar: ", matchedObj.indexOfLastMatchedChar)
                    matchedObj.replaceString()

        # table processing
        for table in doc.tables:
            for  r in table.rows:
                for cell in r.cells:
                    for para in cell.paragraphs:
                        if old_str in para.text:
                            matchedIndices = findMatchedIndices(para.text, 0, old_str)
                            for i_matchIndex, matchIndex in enumerate(matchedIndices):
                                # make the instance of matchedRunObj
                                diff = len(new_str) - len(old_str)
                                updatedMatchedIndex = i_matchIndex * diff + matchIndex
                                matchedObj = MatchedRunObj(para, updatedMatchedIndex, old_str, new_str)
                                matchedObj.findRelatedRunIndex()
                                print("indexOfFirstMatchedChar: ", matchedObj.indexOfFirstMatchedChar)
                                print("indexOfLastMatchedChar: ", matchedObj.indexOfLastMatchedChar)
                                matchedObj.replaceString()
        #heading processing


        # save file in here
        doc.save('.\sample\ex2.docx')
    except AttributeError as e:
        errorObj = ErrorHandler(path, e)
        errorObj.showError()
        return errorObj


def show(path):
    doc = open_docx(path)
    for p in doc.paragraphs:
        for irun, run in enumerate(p.runs):
            print("===:" + str(irun) + ":=================")
            print("|" + run.text + "|")
            print("====================")

def showParagraph(path):
    doc = open_docx(path)
    for ip , p in enumerate(doc.paragraphs):
        print(">>>:" + str(ip) + ":>>>>>>>>>>>>>")
        print(p.text)
        print("<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<")


def travesalTable(path):
    doc = open_docx(path)
    for i_table, table in enumerate(doc.tables):
        for i_row, r in enumerate(table.rows):
            for cell in r.cells:
                for p in cell.paragraphs:
                    print("hello")

def readHeader(path):
    doc = open_docx(path)
    for section in doc.sections:
        print(dir(section))
        # for header in section.header:
        #     print(dir(header))
oStr = "Lara"
nStr = "Pham Thi Thu Trang"
# show(DEFAULT_PATH)
# replace_docx(DEFAULT_PATH, oStr, nStr)
# showParagraph(DEFAULT_PATH)

# travesalTable(DEFAULT_PATH)
# readHeader(DEFAULT_PATH)